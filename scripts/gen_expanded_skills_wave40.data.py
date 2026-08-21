SKILLS = [
    {
        "name": "ai-for-science-communication",
        "title": "AI for Science Communication",
        "description": "Use generative AI to turn scientific findings into clear, audience-tailored public communications—such as plain-language summaries and multimedia explainers—while preserving accuracy.",
        "devin_body": r'''## When to use

You need to translate technical scientific findings into accessible, engaging formats for the public, patients, educators, or policymakers while preserving accuracy.

## Usage

- **Rewrite abstracts and papers into plain-language summaries for non-expert, patient, and public reading levels.**
- **Tune tone, length, and examples for policymakers, journalists, educators, patients, and social media.**
- **Use story structure, metaphors, and relatable examples while preserving uncertainty and avoiding overclaiming.**
- **Combine text, audio, slides, and visuals into accessible, multimodal explainers.**
- **Make every generated claim traceable to the source paper and disclose AI assistance.**

## Steps

1. Identify the target audience, channel, and reading level for the science message.
2. Extract and verify key claims, uncertainties, and source evidence from the original paper or dataset.
3. Generate a plain-language or narrative draft with an LLM prompted for the specific audience and format.
4. Enrich the draft with analogies, visuals, or multimedia while preserving scientific nuance.
5. Fact-check every claim against the source, cite evidence, and disclose AI assistance.
6. Pilot test with a sample audience and refine for comprehension, trust, and accessibility.

## Code pattern

```python
import textstat

# Example: post-process a plain-language summary for target reading level
summary = "CRISPR edits DNA to treat disease."
print("Flesch-Kincaid grade:", textstat.flesch_kincaid_grade(summary))
```

## Tuning notes

- Target a specific reading level (e.g., Flesch-Kincaid 8-10 for general public).
- Preserve hedges and uncertainty (e.g., "suggests," "may," "in this sample").
- Always have a domain expert review AI-generated summaries before publication.
- Disclose AI assistance and maintain transparency about the source material.

## Verification

1. Generate a plain-language summary from a paper and compare its reading level to a human-written version.
2. Fact-check every generated claim against the original source.
3. Pilot the summary with a small non-expert audience and collect comprehension and trust metrics.
''',
        "references": [
            "https://doi.org/10.1177/10755470251411176",
            "https://doi.org/10.1093/pnasnexus/pgae387",
            "https://doi.org/10.48550/arxiv.2308.16377",
            "https://journals.plos.org/plosone/article?id=10.1371/journal.pone.0342852",
        ],
    },
    {
        "name": "ai-for-research-communication",
        "title": "AI for Research Communication",
        "description": "Use LLMs to draft, refine and translate academic manuscripts, cover letters, responses to reviewers and interdisciplinary summaries while grounding every claim in verified sources.",
        "devin_body": r'''## When to use

You are writing or refining academic manuscripts, abstracts, cover letters, response-to-reviewers, or interdisciplinary summaries of research findings.

## Usage

- **Draft and revise IMRaD sections, abstracts, and cover letters that follow journal guidelines.**
- **Adjust formality, clarity, and field-specific conventions for the target venue.**
- **Ground drafts in uploaded PDFs and verified bibliographies, not invented DOIs.**
- **Detect accidental plagiarism, AI-patterned phrasing, and citation errors before submission.**
- **Reframe findings for readers in adjacent fields and for broader impact statements.**

## Steps

1. Assemble source materials: paper, outline, reviewer comments, target journal guidelines, and reference library.
2. Generate a structured first draft of the section (abstract, cover letter, response to reviewers) using the sources as context.
3. Refine tone, length, and terminology to match the journal or correspondence style.
4. Verify every citation, DOI, statistic, and claim against the original sources.
5. Run integrity and style checks, then compare the draft to the original human version.
6. Finalize with author edits and maintain a record of AI involvement for transparency.

## Code pattern

```python
from transformers import pipeline

# Example: condense a long methods section into an abstract-sized summary
summarizer = pipeline("summarization", model="facebook/bart-large-cnn")
short = summarizer(long_methods, max_length=120, min_length=30, do_sample=False)
```

## Tuning notes

- Feed the model your own sources, outline, and reviewer comments to keep outputs grounded.
- Avoid asking the model to invent citations; verify every DOI and page number.
- Maintain author voice by editing generated drafts rather than publishing them raw.
- Use dedicated tools (e.g., Paperpal, Elicit, Semantic Scholar) for citation-aware writing.

## Verification

1. Draft an abstract from a full paper and compare it to the original for accuracy and style.
2. Run a reference check to confirm every generated citation exists and supports its claim.
3. Have a colleague compare an LLM-edited response-to-reviewers to a human-only version.
''',
        "references": [
            "https://peelback.ai/",
            "https://paperpal.com/",
            "https://sciencecast.org/",
            "https://github.com/microsoft/ResearchStudio",
            "https://aclanthology.org/2025.aisd-main.4.pdf",
        ],
    },
    {
        "name": "ai-for-public-engagement",
        "title": "AI for Public Engagement",
        "description": "Use conversational AI and citizen-science chatbots to make public consultations, science festivals and participatory research more inclusive and scalable.",
        "devin_body": r'''## When to use

You are running public consultations, citizen-science projects, science-festival chatbots, or community outreach and want to make engagement more inclusive and scalable.

## Usage

- **Collect, analyze, and respond to public questions, concerns, and ideas.**
- **Deploy chatbots and voice agents that answer science questions and guide participation.**
- **Onboard volunteers, validate submissions, and provide real-time feedback in citizen-science projects.**
- **Support community voice and agency without replacing human decision-making.**
- **Disclose AI involvement, support multilingual interactions, and protect privacy.**

## Steps

1. Co-design engagement goals, prompts, and fallback rules with community stakeholders.
2. Build a retrieval-augmented chatbot grounded in vetted FAQs, papers, and institutional sources.
3. Deploy the agent on accessible channels (web, SMS, voice, event kiosks) in relevant languages.
4. Collect questions and feedback, then extract themes using topic modeling or LLM summarization.
5. Validate chatbot answers against sources and monitor for bias, misinformation, and escalation needs.
6. Iterate with participants and report how input influenced research or policy outcomes.

## Code pattern

```python
from collections import Counter

# Example: simple theme extraction from public consultation comments
def extract_theme(comment):
    # In practice, use an NER or topic model
    return comment.split(":")[0]

themes = [extract_theme(c) for c in comments]
print(Counter(themes).most_common(10))
```

## Tuning notes

- Co-design prompts with community stakeholders, not just technical staff.
- Use retrieval-augmented generation to ground chatbot answers in vetted FAQs and sources.
- Monitor for bias, misinformation, and over-reliance on AI in sensitive discussions.
- Ensure data ownership and consent, especially for youth and marginalized groups.

## Verification

1. Deploy a chatbot at a public event and log question types, answer accuracy, and escalation rates.
2. Analyze a corpus of consultation comments and compare AI-extracted themes to human coding.
3. Measure changes in volunteer retention and data quality when adding an LLM onboarding assistant.
''',
        "references": [
            "https://doi.org/10.1057/s41599-026-06594-5",
            "https://publichealth.jmir.org/2025/1/e65699",
            "https://doi.org/10.1038/s41893-024-01489-2",
            "https://theoryandpractice.citizenscienceassociation.org/articles/10.5334/cstp.812",
        ],
    },
    {
        "name": "ai-for-policy-briefs",
        "title": "AI for Policy Briefs",
        "description": "Convert scientific evidence and legislative text into concise, decision-ready policy briefs and impact analyses for government agencies, regulators and advocacy groups.",
        "devin_body": r'''## When to use

You need to turn a scientific paper, a body of evidence, or a legislative document into a short, decision-ready policy brief for government, agencies, or advocacy groups.

## Usage

- **Assemble problem, evidence, policy options, recommendations, and implications.**
- **Reframe technical findings into actionable, audience-specific guidance.**
- **Map who is affected, how, and what trade-offs exist.**
- **Combine multiple studies while tracking source credibility and recency.**
- **Ensure briefs do not invent statistics, legal clauses, or citations.**

## Steps

1. Define the decision-maker, policy question, and the brief's length and format.
2. Gather and appraise evidence from peer-reviewed research, official legislation, and government data.
3. Synthesize key findings into problem, options, and recommendation statements with citations.
4. Quantify or map stakeholder impacts, costs, and trade-offs for each policy option.
5. Verify every statistic, legal clause, and citation against its original source.
6. Submit the brief for expert review and test clarity with a policy professional or target reader.

## Code pattern

```python
from sklearn.feature_extraction.text import TfidfVectorizer

# Example: identify key themes across a set of policy documents
vectorizer = TfidfVectorizer(max_features=20, stop_words="english")
X = vectorizer.fit_transform(policy_docs)
print(vectorizer.get_feature_names_out())
```

## Tuning notes

- Tailor length and tone to the specific decision-maker and meeting context.
- Use verified sources (peer-reviewed research, official legislation, government data).
- Lead with the recommendation; support it with concise evidence and trade-offs.
- Have a policy expert review the brief before it reaches decision-makers.

## Verification

1. Generate a one-page brief from a scientific paper and compare it to a human-written brief.
2. Verify that every statistic and citation in the brief exists and is accurately represented.
3. Ask a policy professional to rate clarity, relevance, and actionability.
''',
        "references": [
            "https://doi.org/10.48550/arxiv.2509.21493",
            "https://openreview.net/forum?id=S6gJESWNSX",
            "https://doi.org/10.1038/d41586-023-02999-3",
            "https://algorithms.dk/responsible-use-of-ai-in-scientific-advice/",
        ],
    },
    {
        "name": "ai-for-white-papers",
        "title": "AI for White Papers",
        "description": "Author long-form, evidence-based white papers and thought-leadership documents that define problems, survey evidence and present solutions while establishing credibility.",
        "devin_body": r'''## When to use

You are producing a B2B or technical white paper that defines a problem, surveys evidence, presents a solution, and establishes thought leadership.

## Usage

- **State the problem, quantify the cost, and present an evidence-based approach.**
- **Write the executive summary last, place it first, and make it self-contained.**
- **Support claims with benchmarks, peer-reviewed studies, and real customer outcomes.**
- **Include practical guidance on cost, benefits, and adoption.**
- **Maintain a consistent, professional tone and visual format.**

## Steps

1. Define the audience, objective, and key evidence sources before drafting.
2. Research and synthesize industry data, benchmarks, peer-reviewed studies, and customer case studies.
3. Build a structured outline with problem, evidence, solution, implementation, and ROI sections.
4. Draft the body section by section, feeding the model verified sources and avoiding invented citations.
5. Write the executive summary as a self-contained synthesis of the full paper.
6. Fact-check all claims, align with brand voice and design, and export to editable formats.

## Code pattern

```python
from jinja2 import Template

# Example: fill a white-paper outline template from structured evidence
template = Template(open("whitepaper_template.md").read())
doc = template.render(
    title="Edge AI for Manufacturing",
    problem="High latency and cloud costs",
    evidence=evidence_list,
)
```

## Tuning notes

- Decouple evidence retrieval from drafting; never let the model invent sources.
- Define audience, length, and sections before generating text.
- Use a second model or human reviewer to challenge unsupported claims.
- Export to editable formats (DOCX, PDF) with consistent styles and branding.

## Verification

1. Produce an executive summary and confirm it accurately reflects the full paper.
2. Spot-check every statistic and citation against its original source.
3. Compare the AI-assisted white paper to a prior human-written one for tone and structure.
''',
        "references": [
            "https://journals.sagepub.com/doi/10.1177/00472816251332208",
            "https://doi.org/10.1007/978-981-95-4632-9_10",
            "https://specswriter.com/blog/ai_white_papers_how_to_write_one_people_actually_finish.php",
            "https://www.qwe.edu.pl/tutorial/how-to-use-ai-to-write-white-papers/",
        ],
    },
    {
        "name": "ai-for-technical-blogs",
        "title": "AI for Technical Blogs",
        "description": "Plan, draft, SEO-optimize and review technical blog posts and tutorials that combine code, narrative and practical guidance for developer audiences.",
        "devin_body": r'''## When to use

You are creating tutorials, engineering deep-dives, API explainers, or product announcements for a technical audience.

## Usage

- **Combine motivation, concept, code, and outcome in a coherent arc.**
- **Include runnable snippets, output, and common pitfalls.**
- **Structure headings, metadata, and keywords for search and social sharing.**
- **Adjust depth for beginners, practitioners, or experts.**
- **Have subject-matter experts validate accuracy before publishing.**

## Steps

1. Choose a target keyword, audience level, and measurable goal (traffic, engagement, tutorial completion).
2. Create an outline with clear headings, code examples, and a narrative arc from problem to solution.
3. Draft the post with the model, feeding it existing code, docs, and actual output.
4. Run every code snippet in a clean environment and capture real results and error cases.
5. Optimize headings, meta description, and internal links for search without keyword stuffing.
6. Peer-review for technical accuracy, edit for voice, and publish with diagrams and alt text.

## Code pattern

```python
import frontmatter

# Example: write a structured Markdown post with YAML frontmatter
post = frontmatter.Post(
    content=draft_md,
    title="Getting Started with LoRA Fine-Tuning",
    tags=["machine-learning", "fine-tuning"],
    author="Your Name",
)
with open("post.md", "w") as f:
    f.write(frontmatter.dumps(post))
```

## Tuning notes

- Start with a strong outline and code examples; let the model expand, not invent.
- Run every code snippet in the target environment and include actual output.
- Keep the author's voice by editing extensively rather than publishing raw output.
- Add diagrams and alt text to make the post accessible and shareable.

## Verification

1. Publish a draft and compare reader engagement to a manually written baseline.
2. Run all code examples in a clean environment and confirm they execute correctly.
3. Ask a peer to rate technical accuracy, clarity, and usefulness.
''',
        "references": [
            "https://aclanthology.org/2026.findings-acl.296.pdf",
            "https://dev.to/neeraj_ciju/building-vtob-turning-youtube-videos-into-technical-blog-posts-with-a-multi-stage-ai-pipeline-1mng",
            "https://techwriting.pro/",
            "https://www.silverthreadlabs.com/products/bloggen",
            "https://github.com/SurajBhar/deep-blog-agent",
        ],
    },
    {
        "name": "ai-for-open-science",
        "title": "AI for Open Science",
        "description": "Make research reproducible and auditable by automating literature review, code execution, provenance tracking and FAIR data packaging with AI agents.",
        "devin_body": r'''## When to use

You want to make a research project open, reproducible, and auditable by automating literature review, code execution, provenance tracking, and FAIR data packaging.

## Usage

- **Share open data, code, protocols, preprints, and transparent methods.**
- **Build containerized, documented, and versioned replication packages.**
- **Record the origin and transformation of every dataset, figure, and model.**
- **Deploy agents that search literature, run experiments, and write reports with traceability.**
- **Make data Findable, Accessible, Interoperable, and Reusable.**

## Steps

1. Organize the project with versioned data, code, environment files, and a clear README.
2. Use an agent or script to search literature, extract methods, and draft reproducible analysis notebooks.
3. Track provenance with content hashes, container definitions, and RO-Crate or PROV-O metadata.
4. Run the analysis end-to-end and compare outputs to expected values and original data.
5. Package results as a replication archive with figures linked to the scripts that produced them.
6. Share the package under an open license and attempt an independent reproduction by a colleague.

## Code pattern

```python
import hashlib

# Example: create a content hash to track data provenance
with open("data.csv", "rb") as f:
    digest = hashlib.sha256(f.read()).hexdigest()
print("data.csv sha256:", digest)
```

## Tuning notes

- Prefer open-weight or local models when handling sensitive research data.
- Version data, code, and environment definitions together.
- Document every assumption, parameter, and random seed.
- Have an independent run attempt to reproduce the key results.

## Verification

1. Hand the project to a colleague and ask them to reproduce the main result from the README.
2. Compare AI-generated analysis outputs to the original data and code.
3. Check that every figure can be traced back to the script and dataset that produced it.
''',
        "references": [
            "https://arxiv.org/abs/2412.17859",
            "https://github.com/synthetic-sciences/openscience",
            "https://github.com/opencodon/opencodon",
            "https://reproai.org/",
            "https://arxiv.org/abs/2409.11363",
        ],
    },
    {
        "name": "ai-for-data-journalism",
        "title": "AI for Data Journalism",
        "description": "Find, verify and visualize stories in public datasets, documents and real-time data streams to produce data-driven investigative reporting.",
        "devin_body": r'''## When to use

You are investigating public datasets, leaked documents, FOIA releases, or real-time data streams and need to find and verify stories at speed.

## Usage

- **Algorithmically discover, monitor, and verify stories.**
- **Read CSV, JSON, PDF tables, and APIs with reproducible scripts.**
- **Identify people, organizations, and outliers in large corpora.**
- **Link every number, quote, and chart to the underlying source.**
- **Use full-text search, named-entity recognition, and cross-document linking.**

## Steps

1. Acquire public datasets, FOIA releases, or scraped documents and document the source and date.
2. Parse, clean, and join tables with reproducible scripts, tracking each transformation.
3. Use statistics, LLMs, or entity extraction to find anomalies, trends, and story leads.
4. Build charts and interactive graphics and ensure every value matches the source table.
5. Fact-check generated claims by locating the exact row, sentence, or document they came from.
6. Publish the methodology and data alongside the story for transparency and reproducibility.

## Code pattern

```python
import pandas as pd
import altair as alt

# Example: explore a dataset and export an interactive chart
df = pd.read_csv("public_spending.csv")
chart = alt.Chart(df).mark_bar().encode(
    x="department:N",
    y="amount:Q",
)
chart.save("spending_chart.html")
```

## Tuning notes

- Clean and document every data transformation; share the analysis notebook.
- Cross-check surprising findings with the source agency or a domain expert.
- Avoid ecological fallacy; report uncertainty and sample limits.
- Protect sources and personally identifiable information.

## Verification

1. Replicate a published data story from raw data and compare the headline numbers.
2. Generate a chart and manually verify a subset of values against the source table.
3. Fact-check a generated claim by locating the exact sentence or row it came from.
''',
        "references": [
            "https://doi.org/10.48550/arxiv.2606.11176",
            "https://github.com/icij/datashare/",
            "https://datashare.icij.org/",
            "https://www.mdpi.com/2227-7080/10/3/68",
            "https://arxiv.org/abs/2409.07286",
        ],
    },
    {
        "name": "ai-for-visual-communication",
        "title": "AI for Visual Communication",
        "description": "Create and refine posters, slides, pitch decks and social-media assets using diffusion models, layout tools and human-in-the-loop design.",
        "devin_body": r'''## When to use

You need to create presentations, pitch decks, posters, social media assets, or brand visuals that communicate complex ideas clearly and consistently.

## Usage

- **Guide diffusion and layout models with precise, style-aware prompts.**
- **Balance text, images, whitespace, and hierarchy for the target medium.**
- **Enforce colors, fonts, and templates across generated variants.**
- **Constrain generation to layouts, sketches, or existing assets.**
- **Have designers refine AI drafts for accuracy, accessibility, and taste.**

## Steps

1. Define the message, audience, medium, and brand constraints before generating.
2. Generate several visual drafts using style-aware prompts, sketches, or ControlNet constraints.
3. Select the strongest draft and check brand alignment, color contrast, and visual hierarchy.
4. Refine text, labels, and data representations to avoid misrepresentation or bias.
5. Export to editable formats (SVG, PPTX) so designers can finalize and review.
6. Test comprehension and recall with the target audience and review for artifacts and rights issues.

## Code pattern

```python
from PIL import Image, ImageDraw, ImageFont

# Example: build a simple poster canvas in Python
canvas = Image.new("RGB", (1200, 1600), "white")
draw = ImageDraw.Draw(canvas)
draw.text((60, 60), "Research Highlights", fill="black")
canvas.save("poster_draft.png")
```

## Tuning notes

- Generate several variants and select the one that best matches the message.
- Check for visual artifacts, unintended bias, and misrepresentation of data.
- Export editable formats (SVG, PPTX) so designers can refine the output.
- Ensure accessible color contrast and include alt text for generated images.

## Verification

1. Produce a poster or slide deck and compare it to existing brand guidelines.
2. Test the visual with the target audience and measure comprehension and recall.
3. Review generated images for artifacts, false labels, and copyright issues.
''',
        "references": [
            "https://www.nature.com/articles/s41598-026-55838-6",
            "https://www.mdpi.com/2313-433X/11/9/289",
            "https://arxiv.org/abs/2604.04192v1",
            "https://doi.org/10.1016/j.heliyon.2024.e40037",
        ],
    },
    {
        "name": "ai-for-infographics",
        "title": "AI for Infographics",
        "description": "Turn reports, data tables and articles into data-rich infographics and visual stories using natural-language prompts and chart-composition tools.",
        "devin_body": r'''## When to use

You want to turn reports, data tables, or articles into shareable infographics, data stories, or social-media explainers.

## Usage

- **Generate metadata, chart code, and layout from natural language or documents.**
- **Combine multiple sub-charts (bar, line, pie, maps) into a coherent layout.**
- **Ensure values, labels, and visual proportions match the source data.**
- **Guide the eye, use alt text, and maintain color-blind safe palettes.**
- **Use benchmarks like IGENBENCH to assess reliability of generated infographics.**

## Steps

1. Extract the key data, insights, and narrative from the source document or table.
2. Choose the chart types and layout that best communicate the story for the target channel.
3. Generate chart code and metadata, then verify every value and label against the source.
4. Compose sub-charts into an on-brand layout with clear visual hierarchy and alt text.
5. Evaluate the infographic for data faithfulness with a rubric or benchmark like IGENBENCH.
6. Test engagement and comprehension with both data experts and general readers.

## Code pattern

```python
import matplotlib.pyplot as plt

# Example: generate a chart component for later infographic assembly
categories = ["A", "B", "C"]
values = [12, 19, 8]
fig, ax = plt.subplots(figsize=(4, 3))
ax.bar(categories, values)
plt.savefig("chart_component.png")
```

## Tuning notes

- Verify every number and label against the source data table.
- Keep brand colors, fonts, and layout grids consistent.
- Avoid chart junk; prioritize the story over decoration.
- Test generated infographics with both data experts and general readers.

## Verification

1. Generate an infographic from a small table and verify every value and label.
2. Evaluate the output on a reliability benchmark or with a rubric for data faithfulness.
3. Compare engagement and comprehension between the infographic and the original table.
''',
        "references": [
            "https://aclanthology.org/2025.acl-long.1003.pdf",
            "https://aclanthology.org/2026.acl-long.1713.pdf",
            "https://aclanthology.org/anthology-files/anthology-files/pdf/acl/2023.acl-demo.11.pdf",
            "https://arxiv.org/abs/2401.13245",
            "https://arxiv.org/abs/2505.18668v3",
        ],
    },
    {
        "name": "ai-for-document-design",
        "title": "AI for Document Design",
        "description": "Generate consistent, on-brand reports and certificates, proposals and invoices from structured data using templates, typography rules and multi-format rendering.",
        "devin_body": r'''## When to use

You need to produce many reports, certificates, proposals, invoices, or policy briefs from structured data while keeping layouts consistent and on-brand.

## Usage

- **Produce structured page layouts from content and design constraints.**
- **Create reusable templates with dynamic fields for text, tables, and images.**
- **Map CSV, JSON, or database records into document fields.**
- **Output PDF, DOCX, PPTX, or HTML from a single source of truth.**
- **Choose readable fonts, spacing, color contrast, and tagged PDFs.**

## Steps

1. Design and validate a template with brand fonts, colors, margins, and dynamic field placeholders.
2. Connect the template to a CSV, JSON, or database source and map fields to content areas.
3. Add conditional logic for optional sections and repeating rows for tables and lists.
4. Render a batch of documents and inspect page breaks, headers, footers, and formatting.
5. Validate that dynamic fields are bound and no placeholder text remains.
6. Test output across PDF, DOCX, PPTX, or HTML and add accessibility tags where required.

## Code pattern

```python
from jinja2 import Template
from docx import Document

# Example: render a Word report from a template and data
template = Template(open("report_template.md").read())
rendered = template.render(data=records[0])
doc = Document()
doc.add_heading("Report", level=1)
doc.add_paragraph(rendered)
doc.save("report.docx")
```

## Tuning notes

- Design the template once and validate it before batch generation.
- Use conditional logic for optional sections and repeating rows for tables.
- Test page breaks, headers, and footers across edge cases.
- Add PDF/UA or DOCX accessibility tags where required.

## Verification

1. Generate 100 documents from a CSV and visually inspect for formatting consistency.
2. Compare a generated document to a manually produced reference for layout fidelity.
3. Validate that dynamic fields are correctly bound and no placeholder text remains.
''',
        "references": [
            "https://arxiv.org/abs/2510.26213v2",
            "https://doi.org/10.48550/arxiv.2303.10787",
            "https://dl.acm.org/doi/10.1007/978-3-031-41676-7_21",
            "https://www.box.com/docgen",
            "https://imaginepdf.com/",
        ],
    },
    {
        "name": "ai-for-knowledge-design",
        "title": "AI for Knowledge Design",
        "description": "Build knowledge architectures, taxonomies, ontologies and agent-facing knowledge layers that help humans and AI navigate organizational information.",
        "devin_body": r'''## When to use

You are building a knowledge base, wiki, documentation site, knowledge graph, or agent-facing memory system for a team or organization.

## Usage

- **Structure content so both humans and AI agents can navigate it.**
- **Define concepts, relations, and inference rules for a domain.**
- **Connect entities and facts for search, reasoning, and recommendation.**
- **Choose between retrieval at query time and curated, versioned knowledge stores.**
- **Version-control knowledge in Markdown, YAML, or structured schemas.**

## Steps

1. Elicit the questions users and agents need to answer and inventory existing knowledge sources.
2. Define a schema, taxonomy, or ontology with concepts, relations, and source attribution.
3. Extract entities and relationships from documents and curate them with human reviewers.
4. Build a versioned knowledge graph or persistent knowledge layer and link it to RAG or agent tools.
5. Expose the knowledge through both human-readable pages and machine-readable APIs or MCP interfaces.
6. Monitor freshness, coverage, and retrieval quality against representative queries and a reference set.

## Code pattern

```python
import networkx as nx

# Example: build a small knowledge graph from extracted relationships
G = nx.DiGraph()
G.add_node("LoRA", type="technique")
G.add_node("Fine-tuning", type="task")
G.add_edge("LoRA", "Fine-tuning", relation="used_for")
print(nx.shortest_path(G, source="LoRA", target="Fine-tuning"))
```

## Tuning notes

- Start with the questions users and agents need to answer, then design the schema.
- Keep source attribution and freshness metadata on every knowledge unit.
- Use human-in-the-loop curation to avoid compounding AI-generated errors.
- Plan for both human-readable pages and machine-readable APIs/MCP interfaces.

## Verification

1. Build a knowledge graph for a domain and answer a set of representative queries.
2. Check coverage and freshness against a human-curated reference set.
3. Test that an agent can correctly retrieve and cite knowledge in downstream tasks.
''',
        "references": [
            "https://doi.org/10.1080/09544828.2026.2680617",
            "https://link.springer.com/chapter/10.1007/978-3-031-95901-1_1",
            "https://github.com/cantara/knowledge-context-protocol",
            "https://towardsdatascience.com/designing-a-persistent-knowledge-layer-that-refuses-to-guess/",
            "https://knowledge-as-code.com/",
        ],
    },
]